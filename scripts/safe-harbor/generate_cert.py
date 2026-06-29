#!/usr/bin/env python
"""
Safe Harbor Certification generator.

Builds a SC Code 12-37-220(B)(11)(e) / IRS Rev. Proc. 96-32 "Safe Harbor
Certification" letter (modeled on the GSP TTAC template) plus an Exhibit A
unit-by-unit AMI analysis, from:

  * an AppFolio Rent Roll export (.xlsx)
  * an AppFolio Tenant Demographic / Income Range export (.xlsx)  [corroborating only]
  * an entity config (.json) of CAHP/owner boilerplate (see entity_config.template.json)
  * the baked FY2026 HUD/SC Housing rent limits (hud_limits_fy2026.json)

Qualification is RENT-BASED (the basis the template itself states): each unit's
gross rent is compared to the published Maximum Allowable Gross Rent for its
county + bedroom size at 50% / 60% / 80% AMI. The portfolio is then tested
against BOTH Rev. Proc. 96-32 set-aside scopes and the result is reported:

  * 20/50 scope: >= 20% of units at <= 50% AMI  AND  <= 25% above 80% AMI
  * 40/60 scope: >= 40% of units at <= 60% AMI  AND  <= 25% above 80% AMI

(The "<= 25% above 80%" cap is equivalent to the >= 75% low-income overlay.)

The demographic report is attached/summarized as corroborating evidence only;
it cannot drive per-unit determinations (typically mostly "No Income Reported").

Usage:
  python generate_cert.py \
    --rent-roll  rent_roll-20260611.xlsx \
    --demographic tenant_demographic_06-11-2026-03-15-pm.xlsx \
    --entity     safe-harbor-output/iv-spb-ii.entity.json \
    --out        safe-harbor-output \
    [--limits hud_limits_fy2026.json] [--utility-allowance 0] [--tax-year 2026]
"""
from __future__ import annotations

import argparse
import json
import math
import re
import sys
from datetime import date
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

HERE = Path(__file__).resolve().parent
TEAL = RGBColor(0x0F, 0x6E, 0x6E)   # NewShire brand teal-ish for headings


# --------------------------------------------------------------------------- #
# Rent roll parsing
# --------------------------------------------------------------------------- #
class Unit:
    __slots__ = (
        "prop", "unit", "city", "county", "bedrooms", "baths", "tenant",
        "status", "market_rent", "contract_rent", "occupied", "non_residential",
        "gross_rent", "tier", "ceil50", "ceil60", "ceil80", "notes", "source",
    )

    def __init__(self, **kw):
        for k in self.__slots__:
            setattr(self, k, kw.get(k))
        self.notes = kw.get("notes", [])


CITY_RE = None  # built from limits cityToCounty at runtime


def _num(v):
    """Coerce a rent-roll cell to a float, or None."""
    if v is None or v == "":
        return None
    if isinstance(v, (int, float)):
        return float(v)
    s = str(v).replace(",", "").replace("$", "").strip()
    try:
        return float(s)
    except ValueError:
        return None


def detect_county(address: str, city_to_county: dict) -> tuple[str | None, str | None]:
    """Return (city, county) for an AppFolio property string.

    Priority: (1) the city named in the "<City>, SC <zip>" tail, (2) ZIP prefix,
    (3) a loose scan of the whole address. The tail match only accepts a city that
    sits immediately before the state, so a street name containing a city word
    (e.g. "Old Spartanburg Rd Taylors, SC") keys off "Taylors"/the ZIP, not the
    "Spartanburg" buried in the street."""
    addr = address or ""
    # 1) City in the "<City>, SC <zip>" tail — match only cities that end the text
    #    before the state token, so mid-street city words can't win.
    m = re.search(r"(.*?),\s*(?:S\.?C\.?|South Carolina)\b", addr, re.IGNORECASE)
    if m:
        pre = m.group(1).rstrip()
        best = None
        for city in city_to_county:
            if re.search(r"(?:^|\W)" + re.escape(city) + r"\s*$", pre, re.IGNORECASE):
                if best is None or len(city) > len(best):
                    best = city
        if best:
            return best, city_to_county[best]
    # 2) ZIP prefix fallback.
    m = re.search(r"\b(\d{5})(?:-\d{4})?\b", addr)
    if m:
        z = m.group(1)
        if z.startswith("293"):
            return None, "Spartanburg"
        if z.startswith("296"):
            return None, "Greenville"
    # 3) Last resort: any city name anywhere in the address (legacy behavior).
    for city in sorted(city_to_county, key=len, reverse=True):
        if re.search(r"\b" + re.escape(city) + r"\b", addr, re.IGNORECASE):
            return city, city_to_county[city]
    return None, None


def parse_bedrooms(bdba: str):
    """'3/1.00' -> (3, '1.00'); '--/--' -> (None, None)."""
    if not bdba:
        return None, None
    left, _, right = str(bdba).partition("/")
    left, right = left.strip(), right.strip()
    bd = int(left) if left.isdigit() else None
    return bd, (right or None)


def load_rent_roll(path: Path, city_to_county: dict) -> tuple[list[Unit], str, str]:
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))

    # Find the header row (col A == 'Property'), the exported-on date, and the
    # source LLC / property group (so portfolio filings can tag each unit).
    exported = ""
    source = ""
    header_idx = None
    for i, r in enumerate(rows):
        a = (r[0] or "")
        if isinstance(a, str) and a.startswith("Exported On"):
            exported = a.replace("Exported On:", "").strip()
        if isinstance(a, str) and a.startswith("Property Groups"):
            source = re.sub(r"^Owner-\s*", "", a.split(":", 1)[1].strip()).strip()
        if isinstance(a, str) and a.strip() == "Property":
            header_idx = i
            break
    if header_idx is None:
        raise ValueError("Could not find the rent-roll header row (col A == 'Property').")
    if not source:
        source = path.stem  # fall back to the filename when no group line is present

    units: list[Unit] = []
    for r in rows[header_idx + 1:]:
        a = r[0]
        if a is None or str(a).strip() == "":
            continue
        if str(a).strip().lower() == "total":
            break
        prop = str(a).strip()
        unit_lbl = str(r[1]).strip() if r[1] else ""
        bdba = r[3]
        tenant = str(r[4]).strip() if r[4] else ""
        status = str(r[5]).strip() if r[5] else ""
        market_rent = _num(r[7])
        contract_rent = _num(r[8])

        city, county = detect_county(prop, city_to_county)
        bedrooms, baths = parse_bedrooms(bdba)

        non_res = bool(re.search(r"lot only", prop, re.IGNORECASE))
        occupied = bool(tenant) and not status.lower().startswith("vacant")
        gross = contract_rent if (occupied and contract_rent) else market_rent

        u = Unit(
            prop=prop, unit=unit_lbl, city=city, county=county,
            bedrooms=bedrooms, baths=baths, tenant=tenant, status=status,
            market_rent=market_rent, contract_rent=contract_rent,
            occupied=occupied, non_residential=non_res, gross_rent=gross,
            tier=None, ceil50=None, ceil60=None, ceil80=None, notes=[],
            source=source,
        )
        units.append(u)
    return units, exported, source


# --------------------------------------------------------------------------- #
# Tiering
# --------------------------------------------------------------------------- #
def ceilings_for(county_data: dict, bedrooms: int) -> tuple[int, int, int] | None:
    """Return (ceil50, ceil60, ceil80) for a county+bedroom, or None if unavailable."""
    rl = county_data["rentLimits"]
    bkey = str(min(bedrooms, 4))  # tables cap at 4BR; 5+ uses 4BR (flag separately)
    try:
        c50 = rl["50"][bkey]
        c80 = rl["80"][bkey]
    except KeyError:
        return None
    c60 = math.floor(c50 * 1.2)  # MTSP 60% income = 1.2 x 50%; rent scales the same
    return c50, c60, c80


def classify(units: list[Unit], limits: dict, utility_allowance: float):
    counties = limits["counties"]
    for u in units:
        if u.non_residential:
            u.tier = "non-residential"
            u.notes.append("Lot/land only — excluded from residential unit count.")
            continue
        # Units we can't classify (missing county, bedroom, or rent) default to
        # Market — the conservative outcome (counts against qualification, never for).
        if u.county not in counties:
            u.tier = "market"
            u.notes.append("County not determined from address — counted as Market.")
            continue
        if u.bedrooms is None:
            u.tier = "market"
            u.notes.append("Bedroom count missing (BD/BA '--') — counted as Market.")
            continue
        if u.gross_rent is None:
            u.tier = "market"
            u.notes.append("No rent available to test — counted as Market.")
            continue
        if u.bedrooms > 4:
            u.notes.append(f"{u.bedrooms}BR uses the published 4BR ceiling (tables cap at 4BR).")

        cs = ceilings_for(counties[u.county], u.bedrooms)
        if cs is None:
            u.tier = "market"
            u.notes.append("No rent ceiling published for this county/bedroom — counted as Market.")
            continue
        u.ceil50, u.ceil60, u.ceil80 = cs
        gr = u.gross_rent + (utility_allowance or 0)
        if gr <= u.ceil50:
            u.tier = "le50"
        elif gr <= u.ceil60:
            u.tier = "le60"
        elif gr <= u.ceil80:
            u.tier = "le80"
        else:
            u.tier = "market"


def rollup(units: list[Unit]) -> dict:
    residential = [u for u in units if not u.non_residential]
    classified = [u for u in residential if u.tier in ("le50", "le60", "le80", "market")]
    review = [u for u in residential if u.tier == "review"]
    non_res = [u for u in units if u.non_residential]

    # Cumulative counts (a <=50% unit also satisfies <=60% and <=80%).
    c_le50 = sum(u.tier == "le50" for u in classified)
    c_le60 = c_le50 + sum(u.tier == "le60" for u in classified)
    c_le80 = c_le60 + sum(u.tier == "le80" for u in classified)
    c_market = sum(u.tier == "market" for u in classified)

    denom = len(residential)  # all residential units (review units kept in denom)

    def pct(n):
        return round(100.0 * n / denom, 1) if denom else 0.0

    # Worst case: every review unit fails (counts as market / above 80%).
    # Best case:  every review unit lands in the deepest (<=50%) tier.
    n_review = len(review)
    bounds = {
        "le50": (pct(c_le50), pct(c_le50 + n_review)),
        "le60": (pct(c_le60), pct(c_le60 + n_review)),
        "le80": (pct(c_le80), pct(c_le80 + n_review)),
        "market": (pct(c_market), pct(c_market + n_review)),  # (best, worst)
    }

    return {
        "denom": denom,
        "n_review": n_review,
        "n_non_res": len(non_res),
        "counts": {"le50": c_le50, "le60": c_le60, "le80": c_le80, "market": c_market},
        "pct": {
            "le50": pct(c_le50), "le60": pct(c_le60),
            "le80": pct(c_le80), "market": pct(c_market),
        },
        "bounds": bounds,
        "residential": residential,
        "review": review,
        "non_res": non_res,
    }


def evaluate_scopes(roll: dict) -> dict:
    """Determine which Rev. Proc. 96-32 set-aside scope(s) the portfolio meets.

    When review units exist the answer is bracketed: a scope 'qualifies' only if
    it passes even in the worst case (review units treated as failing)."""
    p = roll["pct"]
    has_review = roll["n_review"] > 0
    b = roll["bounds"]

    def verdict(deep_key, deep_threshold):
        # qualifies needs deep% >= threshold AND market% <= 25
        deep_best, deep_worst = b[deep_key]      # (low estimate, high estimate)
        mkt_best, mkt_worst = b["market"]        # (low estimate, high estimate)
        worst_pass = (deep_best >= deep_threshold) and (mkt_worst <= 25.0)
        best_pass = (deep_worst >= deep_threshold) and (mkt_best <= 25.0)
        if worst_pass:
            status = "QUALIFIES"
        elif not best_pass:
            status = "DOES NOT QUALIFY"
        else:
            status = "PROVISIONAL — depends on unresolved units"
        return {
            "deep_pct": p[deep_key],
            "deep_threshold": deep_threshold,
            "market_pct": p["market"],
            "status": status,
        }

    scope_2050 = verdict("le50", 20.0)
    scope_4060 = verdict("le60", 40.0)

    qualifies = [
        name for name, s in (("20/50", scope_2050), ("40/60", scope_4060))
        if s["status"] == "QUALIFIES"
    ]
    if qualifies:
        # Prefer the deeper-affordability scope (20/50) when both pass.
        chosen = "20/50" if "20/50" in qualifies else qualifies[0]
        headline = f"QUALIFIES under the {chosen} scope"
    elif has_review:
        headline = "PROVISIONAL — resolve the units flagged for review, then re-run"
    else:
        headline = "DOES NOT QUALIFY under either scope as configured"

    return {
        "2050": scope_2050, "4060": scope_4060,
        "qualifies": qualifies, "headline": headline, "chosen": qualifies[0] if qualifies else None,
    }


# --------------------------------------------------------------------------- #
# Demographic (corroborating only)
# --------------------------------------------------------------------------- #
def load_demographic(path: Path | None) -> list[tuple[str, str]]:
    if not path:
        return []
    wb = openpyxl.load_workbook(path, data_only=True)
    ws = wb.active
    out = []
    for r in ws.iter_rows(values_only=True):
        if r[0] in (None, "") or str(r[0]).strip().lower().startswith("yearly"):
            continue
        out.append((str(r[0]).strip(), str(r[1]).strip() if r[1] is not None else ""))
    return out


# --------------------------------------------------------------------------- #
# Output: Exhibit A (.xlsx)
# --------------------------------------------------------------------------- #
TIER_LABEL = {
    "le50": "<=50% AMI (very low-income)",
    "le60": "<=60% AMI",
    "le80": "<=80% AMI (low-income)",
    "market": "Market (>80% AMI)",
    "review": "** NEEDS REVIEW **",
    "non-residential": "Non-residential (excluded)",
}


def write_exhibit_a(units: list[Unit], roll: dict, scopes: dict, out_path: Path,
                    entity: dict, utility_allowance: float, group: dict | None = None):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Exhibit A - Unit AMI"
    headers = ["Source LLC", "Property", "Unit", "County", "BR", "Tenant", "Status",
               "Market Rent", "Contract Rent", "Gross Rent Tested",
               "50% Ceil", "60% Ceil", "80% Ceil", "AMI Tier", "Notes"]
    ws.append(headers)
    for c in ws[1]:
        c.font = openpyxl.styles.Font(bold=True)
    for u in sorted(units, key=lambda x: (x.source or "", x.prop)):
        ws.append([
            u.source, u.prop, u.unit, u.county or "?",
            (u.bedrooms if u.bedrooms is not None else "--"),
            u.tenant, u.status, u.market_rent, u.contract_rent,
            (None if u.gross_rent is None else u.gross_rent + (utility_allowance or 0)),
            u.ceil50, u.ceil60, u.ceil80,
            TIER_LABEL.get(u.tier, u.tier or "?"),
            "; ".join(u.notes),
        ])
    # widths
    for col, w in zip("ABCDEFGHIJKLMNO",
                      [22, 34, 14, 12, 5, 22, 16, 12, 13, 16, 9, 9, 9, 26, 40]):
        ws.column_dimensions[col].width = w

    # Summary sheet
    s = wb.create_sheet("Summary")
    p, cnt = roll["pct"], roll["counts"]
    s.append(["Safe Harbor Summary", entity["company"]["legalName"]])
    s.append([])
    s.append(["Total residential units (denominator)", roll["denom"]])
    s.append(["Units needing review", roll["n_review"]])
    s.append(["Non-residential units excluded", roll["n_non_res"]])
    s.append([])
    s.append(["AMI Tier (cumulative)", "Units", "% of Total", "Required", "Result"])
    s.append(["Low-Income (<=80% AMI)", cnt["le80"], f"{p['le80']}%", ">=75%",
              "PASS" if p["le80"] >= 75 else "FAIL"])
    s.append(["  <=60% AMI", cnt["le60"], f"{p['le60']}%", ">=40% (40/60 scope)",
              "PASS" if p["le60"] >= 40 else "FAIL"])
    s.append(["  Very Low-Income (<=50% AMI)", cnt["le50"], f"{p['le50']}%",
              ">=20% (20/50 scope)", "PASS" if p["le50"] >= 20 else "FAIL"])
    s.append(["Market (>80% AMI)", cnt["market"], f"{p['market']}%", "<=25%",
              "PASS" if p["market"] <= 25 else "FAIL"])
    s.append([])
    s.append(["20/50 scope", scopes["2050"]["status"]])
    s.append(["40/60 scope", scopes["4060"]["status"]])
    s.append(["DETERMINATION", scopes["headline"]])
    for col, w in zip("ABCDE", [30, 10, 12, 22, 10]):
        s.column_dimensions[col].width = w

    # Per-LLC breakdown sheet for a group filing.
    if group:
        g = wb.create_sheet("Per-LLC")
        g.append([f"Portfolio: {group['name']}"])
        g.append([])
        g.append(["Source LLC", "Units", "<=50%", "<=60%", "<=80%", "Market",
                  "20/50", "40/60"])
        for c in g[3]:
            c.font = openpyxl.styles.Font(bold=True)
        for src in group["sources"]:
            r = group["per_src"][src]
            sc = group["per_scope"][src]
            rp = r["pct"]
            g.append([src, r["denom"], f"{rp['le50']}%", f"{rp['le60']}%",
                      f"{rp['le80']}%", f"{rp['market']}%",
                      sc["2050"]["status"], sc["4060"]["status"]])
        g.append([])
        g.append(["PORTFOLIO TOTAL", roll["denom"], f"{p['le50']}%", f"{p['le60']}%",
                  f"{p['le80']}%", f"{p['market']}%",
                  scopes["2050"]["status"], scopes["4060"]["status"]])
        g[g.max_row][0].font = openpyxl.styles.Font(bold=True)
        for col, w in zip("ABCDEFGH", [24, 7, 8, 8, 8, 8, 24, 24]):
            g.column_dimensions[col].width = w
    wb.save(out_path)


# --------------------------------------------------------------------------- #
# Output: certification letter (.docx)
# --------------------------------------------------------------------------- #
def _blank(val, width=24):
    return val if val else "_" * width


def _h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.color.rgb = TEAL
    r.font.size = Pt(11)
    return p


def write_letter(units, roll, scopes, demo, entity, limits, out_path: Path,
                 rent_roll_date: str, utility_allowance: float, tax_year: int,
                 group: dict | None = None):
    co = entity["company"]
    prop = entity["property"]
    non = entity["nonprofit"]
    filing = entity["filing"]
    p, cnt = roll["pct"], roll["counts"]
    company_name = group["name"] if group else co["legalName"]
    company_type = group["stateType"] if group else co["stateType"]
    # The certifying party is the property management company (as authorized agent
    # for the owner), not the nonprofit. relationshipToOwner ties the signer to the
    # ownership; the signature block itself is left blank and filled at signing.
    cert = entity.get("certification", {})
    relationship = cert.get("relationshipToOwner", "property manager and authorized agent")
    # Nonprofit ownership %/class come from the hub (per entity) — blank if absent.
    _pct = non.get("ownershipPercent")
    pct_disp = f"{_pct}%" if _pct not in (None, "") else "____%"
    cls_disp = non.get("memberClass") or "____"
    if _pct in (None, ""):
        ownership_clause = "holding the membership interest set forth in its operating agreement"
    else:
        ownership_clause = f"holding a {_words_pct(_pct)} ({_pct}%) {cls_disp} Interest"

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("SAFE HARBOR CERTIFICATION")
    rt.bold = True
    rt.font.size = Pt(15)
    rt.font.color.rgb = TEAL
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("SOUTH CAROLINA CODE §12-37-220(B)(11)(e)\nIRS REVENUE PROCEDURE 96-32").bold = True

    # Parameter block
    counties_str = " and ".join(prop["counties"]) + " " + (
        "Counties" if len(prop["counties"]) > 1 else "County")
    msa_str = "; ".join(sorted({limits["counties"][c]["msa"]
                                for c in prop["counties"] if c in limits["counties"]}))
    desc_value = f"{roll['denom']} {prop['description']}"
    if group:
        desc_value += f" held across {len(group['sources'])} single-purpose LLCs (see Section 2)"
    params = [
        ("Property", prop["addressLine"]),
        ("Description", desc_value),
        ("Entity", f"{company_name} ({company_type})"
                   + (" — portfolio / group filing" if group else "")),
        ("Nonprofit Managing Member",
         f"{non['managingMemberName']} (instrumentality of {non['parentName']})"),
        ("Nonprofit Ownership", f"{pct_disp} {cls_disp} Interest"),
        ("Certifying Party", f"Property management company, as {relationship}"),
        ("CAHP EIN", _blank(non.get("parentEin"), 14)),
        ("Entity EIN", _blank(co.get("ein"), 14)),
        ("County", counties_str),
        ("Tax Year", str(tax_year)),
        ("Filing Type", filing["filingType"]),
        ("Certification Date", "_" * 24),
        ("Rent Roll Date", rent_roll_date or "_" * 16),
        ("HUD/SC Housing Limits", f"FY{limits['_meta']['fiscalYear']} "
                                  f"(effective {limits['_meta']['effectiveDate']}); {msa_str}"),
    ]
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Light List Accent 1"
    for k, v in params:
        cells = tbl.add_row().cells
        cells[0].paragraphs[0].add_run(k).bold = True
        cells[1].text = str(v)

    doc.add_paragraph()
    doc.add_paragraph(
        "TO THE SOUTH CAROLINA DEPARTMENT OF REVENUE AND THE "
        + counties_str.upper() + " ASSESSOR:"
    ).runs[0].bold = True
    company_clause = (
        f"{company_name} (the “Company”) and its wholly-owned single-purpose "
        f"subsidiary LLCs listed in Section 2"
        if group else f"{company_name} (the “Company”)"
    )
    doc.add_paragraph(
        f"The undersigned property management company, as {relationship} for "
        f"{company_clause}, hereby certifies under penalty of perjury as follows:"
    )

    _h(doc, "1. Entity Structure and Nonprofit Status.")
    doc.add_paragraph(
        f"{non['parentName']} (“CAHP”) is a South Carolina nonprofit corporation "
        f"exempt from federal income tax under Section 501(c)(3) of the Internal Revenue "
        f"Code. {non['managingMemberName']} is a wholly owned instrumentality of CAHP and "
        f"serves as the managing member of the Company, {ownership_clause}. The Property "
        f"is managed day-to-day by the undersigned property management company, which "
        f"executes this certification as {relationship} for the Company."
    )

    _h(doc, "2. Property Description.")
    if group:
        doc.add_paragraph(
            f"The Company holds the residential rental real property through "
            f"{_words_num(len(group['sources']))} ({len(group['sources'])}) "
            f"{group['subsidiaryDescription']}, together comprising "
            f"{_words_num(roll['denom'])} ({roll['denom']}) residential rental units "
            f"located in {counties_str}, South Carolina. Unit detail accompanies this "
            f"certification in the submitted rent roll."
        )
    else:
        doc.add_paragraph(
            f"The Company owns the residential rental real property described above, "
            f"consisting of {_words_num(roll['denom'])} ({roll['denom']}) "
            f"{prop['description']} located in {counties_str}, South Carolina."
        )

    _h(doc, "3. Safe Harbor Qualification — Revenue Procedure 96-32.")
    doc.add_paragraph(
        "Qualification is established on a rent-restriction basis: a unit counts toward an "
        "AMI tier when its gross rent (contract rent plus any tenant-paid utility allowance) "
        "does not exceed the Maximum Allowable Gross Rent published for the applicable "
        "county, bedroom size, and that AMI tier under the HUD MTSP methodology (FY"
        f"{limits['_meta']['fiscalYear']} limits, effective "
        f"{limits['_meta']['effectiveDate']}). The required percentage of residential units "
        "is so rent-restricted, satisfying the Rev. Proc. 96-32 set-aside as follows:"
    )
    # Determination sentence
    det = doc.add_paragraph()
    det.add_run("Set-Aside Determination: ").bold = True
    det.add_run(scopes["headline"]).bold = True

    # Certify under a single set-aside program: prefer 20/50 (50% AMI), else
    # 40/60 (60% AMI). Show only that program's deep tier + the shared 75%/25%.
    chosen = scopes.get("chosen")
    program_label = (
        "the 20%-at-50%-AMI set-aside (Rev. Proc. 96-32 §3.02)" if chosen == "20/50"
        else "the 40%-at-60%-AMI set-aside (Rev. Proc. 96-32 §3.02)" if chosen == "40/60"
        else None
    )
    if program_label:
        doc.add_paragraph(
            f"The Property is certified under {program_label}; the test for that "
            f"program is shown below."
        )

    # Tier table
    t = doc.add_table(rows=1, cols=5)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, txt in enumerate(["AMI Tier", "Units", "% Total", "Required", "Result"]):
        hdr[i].paragraphs[0].add_run(txt).bold = True
    # Non-cumulative (partition) counts: each unit counted only in its deepest tier.
    rc = {"le50": 0, "le60": 0, "le80": 0, "market": 0}
    for u in units:
        if u.non_residential:
            continue
        if u.tier in rc:
            rc[u.tier] += 1
    denom_r = roll["denom"] or 1

    def pcr(n):
        return round(100.0 * n / denom_r, 1)

    market_row = ("Market (>80% AMI)", rc["market"], pcr(rc["market"]), "≤25%", pcr(rc["market"]) <= 25)
    if chosen == "20/50":
        rows = [
            ("Very Low-Income (≤50% AMI)", rc["le50"], pcr(rc["le50"]), "≥20%", pcr(rc["le50"]) >= 20),
            ("Low-Income (>50% to ≤80% AMI)", rc["le60"] + rc["le80"], pcr(rc["le60"] + rc["le80"]), "", None),
            market_row,
        ]
    elif chosen == "40/60":
        rows = [
            ("Low-Income (≤60% AMI)", rc["le50"] + rc["le60"], pcr(rc["le50"] + rc["le60"]), "≥40%", pcr(rc["le50"] + rc["le60"]) >= 40),
            ("Low-Income (>60% to ≤80% AMI)", rc["le80"], pcr(rc["le80"]), "", None),
            market_row,
        ]
    else:
        rows = [
            ("≤50% AMI", rc["le50"], pcr(rc["le50"]), "", None),
            (">50% to ≤60% AMI", rc["le60"], pcr(rc["le60"]), "", None),
            (">60% to ≤80% AMI", rc["le80"], pcr(rc["le80"]), "", None),
            market_row,
        ]
    for label, n, pct_v, req, ok in rows:
        c = t.add_row().cells
        c[0].text = label
        c[1].text = str(n)
        c[2].text = f"{pct_v}%"
        c[3].text = req
        c[4].text = "" if ok is None else ("PASS" if ok else "FAIL")

    if roll["n_review"]:
        warn = doc.add_paragraph()
        wr = warn.add_run(
            f"⚠ {roll['n_review']} unit(s) could not be auto-classified (missing "
            f"bedroom count, rent, or county) and are listed in Exhibit A under "
            f"“NEEDS REVIEW.” The percentages above treat them conservatively. "
            f"Resolve them and re-run before signing."
        )
        wr.bold = True
        wr.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    _h(doc, "4. Corroborating Resident Income Data.")
    if demo:
        doc.add_paragraph(
            "The Company's resident income-range report corroborates the affordability "
            "profile reflected by the rent restrictions above. Reported distribution of "
            "current residents by annual income range:"
        )
        dt = doc.add_table(rows=1, cols=2)
        dt.style = "Light List Accent 1"
        dt.rows[0].cells[0].paragraphs[0].add_run("Annual Income Range").bold = True
        dt.rows[0].cells[1].paragraphs[0].add_run("% of Current Residents").bold = True
        for rng, pc in demo:
            c = dt.add_row().cells
            c[0].text = rng
            c[1].text = pc
        doc.add_paragraph(
            "This report is provided as supporting evidence only; unit-level qualification "
            "is established by the rent-restriction analysis in Section 3."
        )
    else:
        doc.add_paragraph("[Resident income-range report not provided.]")

    _h(doc, "5. Exemption Request.")
    doc.add_paragraph(
        f"Based on the foregoing, the Company respectfully requests a full exemption from "
        f"ad valorem property taxation under South Carolina Code §12-37-220(B)(11)(e) "
        f"for tax year {tax_year}."
    )

    # ── Certification & signature — simple, fully fillable at signing ──
    owner_ref = (
        f"{company_name} and the subsidiary LLCs listed in Section 2"
        if group else company_name
    )
    doc.add_paragraph()
    doc.add_paragraph(
        f"Certified under penalty of perjury this _____ day of _______________, {tax_year}."
    )
    auth = doc.add_paragraph()
    ar = auth.add_run(
        "This certification is executed by the property management company that manages the "
        f"Property, as {relationship} for {owner_ref}."
    )
    ar.font.size = Pt(9)
    doc.add_paragraph()
    for line in (
        "Signature: ______________________________",
        "Name:      ______________________________",
        "Title:     ______________________________",
        "Company:   ______________________________",
        "Date:      ______________________________",
    ):
        doc.add_paragraph(line)

    # Source footnote
    fp = doc.add_paragraph()
    fr = fp.add_run(
        f"Rent limits source: {limits['_meta']['source']} "
        f"Verify against the live source before filing."
    )
    fr.font.size = Pt(7.5)
    fr.italic = True

    doc.save(out_path)


def _words_pct(n):
    return {1: "one percent", 2: "two percent", 5: "five percent"}.get(n, f"{n} percent")


_ONES = ["zero", "one", "two", "three", "four", "five", "six", "seven", "eight",
         "nine", "ten", "eleven", "twelve", "thirteen", "fourteen", "fifteen",
         "sixteen", "seventeen", "eighteen", "nineteen"]
_TENS = ["", "", "twenty", "thirty", "forty", "fifty", "sixty", "seventy",
         "eighty", "ninety"]


def _words_num(n):
    if n < 20:
        return _ONES[n]
    if n < 100:
        return _TENS[n // 10] + ("-" + _ONES[n % 10] if n % 10 else "")
    return str(n)


# --------------------------------------------------------------------------- #
# Main
# --------------------------------------------------------------------------- #
def slugify(name: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", name).strip("_")


def main(argv=None):
    ap = argparse.ArgumentParser(description=__doc__,
                                 formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--rent-roll", required=True, type=Path, nargs="+",
                    help="One or more AppFolio rent-roll .xlsx files (one per LLC), "
                         "or a directory of them. Multiple distinct LLCs => a group "
                         "(portfolio) filing aggregated into one certification.")
    ap.add_argument("--group-name",
                    help="Parent/portfolio name for a group filing (e.g. "
                         "'IV Fund Global, LLC'). Overrides entity.portfolio.groupName.")
    ap.add_argument("--demographic", type=Path)
    ap.add_argument("--entity", required=True, type=Path)
    ap.add_argument("--limits", type=Path, default=HERE / "hud_limits_fy2026.json")
    ap.add_argument("--out", type=Path, default=Path("safe-harbor-output"))
    ap.add_argument("--utility-allowance", type=float, default=0.0,
                    help="Flat tenant-paid utility allowance added to contract rent "
                         "when computing gross rent (default 0).")
    ap.add_argument("--tax-year", type=int, default=None)
    args = ap.parse_args(argv)

    limits = json.loads(args.limits.read_text(encoding="utf-8"))
    entity = json.loads(args.entity.read_text(encoding="utf-8"))
    tax_year = args.tax_year or entity.get("filing", {}).get("taxYear", date.today().year)

    # Expand directories, load every rent roll, tag each unit with its source LLC.
    roll_files: list[Path] = []
    for p in args.rent_roll:
        if p.is_dir():
            roll_files.extend(sorted(p.glob("*.xlsx")))
        else:
            roll_files.append(p)
    if not roll_files:
        ap.error("No rent-roll .xlsx files found.")

    units: list[Unit] = []
    rr_dates: list[str] = []
    for p in roll_files:
        us, d, _src = load_rent_roll(p, limits["cityToCounty"])
        units.extend(us)
        if d:
            rr_dates.append(d)

    classify(units, limits, args.utility_allowance)
    roll = rollup(units)
    scopes = evaluate_scopes(roll)
    demo = load_demographic(args.demographic)

    # Group / portfolio filing? (more than one distinct source LLC, or flagged).
    distinct_sources = sorted({u.source for u in units if u.source})
    portfolio = entity.get("portfolio") or {}
    group_name = args.group_name or portfolio.get("groupName")
    is_group = bool(portfolio.get("isGroupFiling")) or len(distinct_sources) > 1
    group = None
    if is_group:
        per_src = {
            s: rollup([u for u in units if u.source == s]) for s in distinct_sources
        }
        per_scope = {s: evaluate_scopes(per_src[s]) for s in distinct_sources}
        group = {
            "name": group_name or entity["company"]["legalName"],
            "stateType": portfolio.get("groupStateType",
                                       "South Carolina limited liability company"),
            "subsidiaryDescription": portfolio.get("subsidiaryDescription",
                                                    "wholly-owned single-purpose subsidiary LLCs"),
            "sources": distinct_sources,
            "per_src": per_src,
            "per_scope": per_scope,
        }

    rr_date = "; ".join(dict.fromkeys(rr_dates))  # de-dup, preserve order

    args.out.mkdir(parents=True, exist_ok=True)
    base = slugify(group["name"]) + "_GROUP" if is_group else slugify(entity["company"]["legalName"])
    letter_path = args.out / f"{base}_Safe_Harbor_Certification_TY{tax_year}.docx"
    exhibit_path = args.out / f"{base}_Unit_AMI_Analysis_INTERNAL.xlsx"

    write_exhibit_a(units, roll, scopes, exhibit_path, entity, args.utility_allowance, group)
    write_letter(units, roll, scopes, demo, entity, limits, letter_path,
                 rr_date, args.utility_allowance, tax_year, group)

    # Console summary
    p, cnt = roll["pct"], roll["counts"]
    print("=" * 64)
    if is_group:
        print(f"  GROUP FILING: {group['name']}  |  Tax Year {tax_year}")
        print(f"  {len(distinct_sources)} LLCs aggregated as one portfolio")
    else:
        print(f"  {entity['company']['legalName']}  |  Tax Year {tax_year}")
    print("=" * 64)
    if is_group:
        print("  Per-LLC breakdown (units | <=50% | <=80% | mkt):")
        for s in distinct_sources:
            r = group["per_src"][s]
            print(f"    {s:<22} {r['denom']:>3} | {r['pct']['le50']:>5}% "
                  f"| {r['pct']['le80']:>5}% | {r['pct']['market']:>5}%")
        print("-" * 64)
        print("  AGGREGATED PORTFOLIO:")
    print(f"  Residential units (denominator): {roll['denom']}")
    print(f"  Non-residential excluded:        {roll['n_non_res']}")
    print(f"  Needs review:                    {roll['n_review']}")
    print("-" * 64)
    print(f"  <=50% AMI : {cnt['le50']:>3}  ({p['le50']}%)   [20/50 needs >=20%]")
    print(f"  <=60% AMI : {cnt['le60']:>3}  ({p['le60']}%)   [40/60 needs >=40%]")
    print(f"  <=80% AMI : {cnt['le80']:>3}  ({p['le80']}%)   [needs >=75%]")
    print(f"  Market    : {cnt['market']:>3}  ({p['market']}%)   [needs <=25%]")
    print("-" * 64)
    print(f"  20/50 scope: {scopes['2050']['status']}")
    print(f"  40/60 scope: {scopes['4060']['status']}")
    print(f"  >>> {scopes['headline']}")
    print("=" * 64)
    print(f"  Letter : {letter_path}")
    print(f"  Unit analysis (internal, not for submittal): {exhibit_path}")
    if roll["n_review"]:
        print("\n  Units needing review:")
        for u in roll["review"]:
            print(f"    - {u.prop} [{u.unit}] : {'; '.join(u.notes)}")


if __name__ == "__main__":
    sys.exit(main())
